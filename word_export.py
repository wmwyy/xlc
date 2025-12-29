"""Word文档导出模块 - 消力池计算结果"""

from __future__ import annotations
from datetime import datetime
from typing import Dict, Any


def _require_docx():
    """检查并导入python-docx依赖"""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt
        return Document, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, qn, Cm, Pt
    except Exception as e:
        raise ImportError("缺少依赖：python-docx（请先 pip install python-docx）") from e


def _ensure_docx_suffix(path: str) -> str:
    """确保文件名以.docx结尾"""
    p = str(path)
    return p if p.lower().endswith(".docx") else (p + ".docx")


def _fmt(x, nd: int = 3) -> str:
    """格式化数值显示"""
    try:
        v = float(x)
        if abs(v) >= 1e4 or (abs(v) > 0 and abs(v) < 1e-3):
            return f"{v:.{nd}e}"
        return f"{v:.{nd}f}".rstrip("0").rstrip(".")
    except Exception:
        return str(x)


def _build_doc_base():
    """创建基础文档对象"""
    Document, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, qn, Cm, Pt = _require_docx()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)

    return doc, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, qn, Cm, Pt


def _add_centered_title(doc, title_text: str):
    """添加居中标题"""
    _, WD_ALIGN_PARAGRAPH, _, _, _, Pt = _require_docx()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title_text)
    run.bold = True
    run.font.size = Pt(16)


def _add_section_title(doc, title_text: str):
    """添加章节标题"""
    _, _, _, _, _, Pt = _require_docx()
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    run.bold = True
    run.font.size = Pt(12)


def _add_line(doc, text: str):
    """添加带Unicode格式的文本行"""
    # Unicode下标和上标
    text = text.replace("h''", "h")
    text = text.replace("h'", "h")
    text = text.replace('_c', 'c')
    text = text.replace('_s', 's')
    text = text.replace('_j', 'j')
    text = text.replace('_sj', 'sj')
    text = text.replace('m', 'm')
    text = text.replace('m', 'm')
    
    doc.add_paragraph(text)


def export_energy_basin_to_word(
    results: Dict[str, Any],
    output_path: str,
    project_name: str = "消力池计算",
    input_params: dict = None
) -> str:
    """
    导出消力池计算结果到Word文档

    Args:
        results: 计算结果字典
        output_path: 输出文件路径
        project_name: 工程名称
        input_params: 输入参数字典

    Returns:
        str: 实际保存的文件路径
    """
    output_path = _ensure_docx_suffix(output_path)
    doc, _, _, _, _, Pt = _build_doc_base()

    # 标题
    _add_centered_title(doc, "消力池计算报告")
    _add_centered_title(doc, f"（{project_name}）")
    doc.add_paragraph()

    # 计算依据
    _add_section_title(doc, "一、计算依据")
    _add_line(doc, "规范附录B.1 - 消力池计算")
    _add_line(doc, "计算内容：收缩断面水深hc、下游水深hc、跃后水深ΔZ等")
    doc.add_paragraph()

    # 输入参数
    if input_params:
        _add_section_title(doc, "二、输入参数")
        _add_line(doc, f"σ - 跳跃淹没系数：{_fmt(input_params.get('sigma0', 0))}")
        _add_line(doc, f"α - 动能校正系数：{_fmt(input_params.get('alpha', 0))}")
        _add_line(doc, f"q - 单宽流量：{_fmt(input_params.get('q', 0))} m/s/m")
        _add_line(doc, f"b - 首端宽度：{_fmt(input_params.get('b1', 0))} m")
        _add_line(doc, f"b - 末端宽度：{_fmt(input_params.get('b2', 0))} m")
        _add_line(doc, f"T - 总势能：{_fmt(input_params.get('T0', 0))} m")
        _add_line(doc, f"p - 校正长度参数：{_fmt(input_params.get('p', 0))} m")
        _add_line(doc, f"hs - 出池河床水深：{_fmt(input_params.get('hs', 0))} m")
        _add_line(doc, f"Ls - 斜段水平投影：{_fmt(input_params.get('Ls', 0))} m")
        _add_line(doc, f"β - 水跃长度校正：{_fmt(input_params.get('beta', 0))}")
        _add_line(doc, f"g - 重力加速度：{_fmt(input_params.get('g', 0))} m/s")
        doc.add_paragraph()

    # 计算结果
    _add_section_title(doc, "三、计算结果")

    _add_line(doc, "【水深计算】")
    _add_line(doc, f"收缩断面水深 hc = {_fmt(results.get('hc', 0))} m")
    _add_line(doc, f"下游水深 hc = {_fmt(results.get('hc_double_prime', 0))} m")
    _add_line(doc, f"跃后水深 ΔZ = {_fmt(results.get('delta_Z', 0))} m")
    _add_line(doc, f"池深 d = {_fmt(results.get('d', 0))} m")
    doc.add_paragraph()

    _add_line(doc, "【长度计算】")
    _add_line(doc, f"水跃长度 Lj = {_fmt(results.get('Lj', 0))} m")
    _add_line(doc, f"护坦长度 Lsj = {_fmt(results.get('Lsj', 0))} m")
    doc.add_paragraph()

    _add_line(doc, "【其他参数】")
    if 'v' in results:
        _add_line(doc, f"流速 v = {_fmt(results.get('v', 0))} m/s")
    if 'Fr' in results:
        _add_line(doc, f"Froude数 Fr = {_fmt(results.get('Fr', 0))}")
    doc.add_paragraph()

    # 计算时间
    _add_line(doc, f"计算时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
    doc.add_paragraph()

    # 页脚说明
    p = doc.add_paragraph()
    run = p.add_run("注：本报告由消力池计算器自动生成")
    run.font.size = Pt(9)
    run.italic = True

    # 保存文档
    doc.save(output_path)
    return output_path
